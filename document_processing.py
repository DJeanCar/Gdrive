import io
from google_drive_service import drive_service
from docx import Document
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload

# REPORT TEMPLATE
FILE_ID = '1PRv1MEtpLBMMuwfRpCiLA8nLPrQeBNr3TmsJZGUtvQs'
FINAL_FILE_NAME = 'final_file.docx'
DOCX_MIME_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
PDF_MIME_TYPE = 'application/pdf'

def download_file():
    request = drive_service.files().export_media(fileId=FILE_ID,
                                             mimeType=DOCX_MIME_TYPE)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print("Download %d%%." % int(status.progress() * 100))

    with open(FINAL_FILE_NAME, "wb") as f:
        f.write(fh.getbuffer())


def fill_file_fields(fields):
    document = Document(FINAL_FILE_NAME)

    for paragraph in document.paragraphs:
        inline = paragraph.runs
        for i in range(len(inline)):
            for key in fields.keys():
                inline[i].text = inline[i].text.replace(key, fields[key])

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        for key in fields.keys():
                            inline[i].text = inline[i].text.replace(key, fields[key])

    document.save(FINAL_FILE_NAME)


def upload_file():
    media = MediaFileUpload(FINAL_FILE_NAME, mimetype=DOCX_MIME_TYPE)
    drive_service.files().create(
        body={
            'name': FINAL_FILE_NAME,
        },
        media_body=media,
        fields='id'
    ).execute()
